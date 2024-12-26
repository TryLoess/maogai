import pandas as pd
from docx import Document

doc = Document()

def add_to_docx(data, doc):
    doc.add_heading(f"{data['题号']}：{data['题目']}（{data['难度']}；{data['题型']}）", level=4)

    options = [f"{chr(65 + i)}. {opt}" if not opt[0].isupper() else opt for i, opt in enumerate(data["选项"])]

    # 写入选项
    for option in options:
        doc.add_paragraph(option)

    # 写入正确答案
    doc.add_paragraph(f"正确答案：{data['正确答案']}")

data = pd.read_json(r"E:\python\建模\web_try\tiku.json")
data.apply(lambda x: add_to_docx(x, doc), axis=1)
doc.save("output.docx")
